from docx import Document
from docx.shared import Pt
import io

def create_docx(text, title="Irresistible Guide"):
    """
    Converts a plain text (or simple markdown) string into a Word document.
    Returns a BytesIO object ready for download.
    """
    doc = Document()
    
    # Title
    doc.add_heading(title, 0)
    
    # Add content
    # A simple parser for headers and lists would go here
    # For now, we just dump the text, but we can make it smarter later
    
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
             doc.add_paragraph(line[3:], style='List Number')
        else:
            doc.add_paragraph(line)
            
    # Save to memory
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
