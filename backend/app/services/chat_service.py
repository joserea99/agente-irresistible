"""
Chat Service - Handles AI conversations with Gemini
Migrated from agent_logic.py
"""

from google import genai
from google.genai import types
from typing import List, Dict, Optional
import os

# Import personas
from .personas import PERSONAS

# Global imports for types/standard libs
import re
from io import BytesIO

# Optional docx import with compatibility
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        WD_ALIGN_CENTER = WD_ALIGN_PARAGRAPH.CENTER
    except ImportError:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        WD_ALIGN_CENTER = WD_PARAGRAPH_ALIGNMENT.CENTER
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ python-docx not installed. Export will fail or fallback.")

class ChatService:
    """Service for handling AI chat conversations with Gemini"""
    
    def __init__(self, api_key: Optional[str] = None):
        """Initialize the chat service with Gemini API"""
        self.api_key = api_key or os.environ.get("GOOGLE_API_KEY")
        
        if self.api_key:
            self.client = genai.Client(api_key=self.api_key)
            self.model_name = "gemini-2.0-flash"
        else:
            self.client = None

    
    def get_directors(self) -> List[Dict[str, str]]:
        """Get list of available directors"""
        directors = []
        for key in PERSONAS.keys():
            # Create a simple ID from the name (e.g., "Pastor Principal" -> "pastor_principal")
            simple_id = key.lower().replace(" ", "_").replace("(", "").replace(")", "")
            directors.append({
                "id": simple_id,
                "name": key,
                "key": key
            })
        return directors
    
    def generate_response(
        self,
        user_input: str,
        history: List[Dict[str, str]] = [],
        director: str = "Programación de Servicio",
        rag_context: Optional[str] = None
    ) -> str:
        """
        Generate AI response
        """
        if not self.client:
            return "⚠️ **Error:** No Google Gemini API Key configured."
        
        # Get system prompt for selected director
        system_prompt = PERSONAS.get(director, PERSONAS["Programación de Servicio"])
        
        # Add RAG context if available
        if rag_context:
            system_prompt += f"""
            
## RELEVANT CONTEXT FROM KNOWLEDGE BASE:
The following is information retrieved from the church's knowledge base that may be relevant to the user's question:

{rag_context}

Use this context to provide more specific and accurate answers. If the context doesn't apply, use your general knowledge.
"""
        
        # Build messages for google-genai
        # It expects a list of Content objects or properly formatted config
        # We'll use the generate_content API which usually takes system instruction in config
        
        contents = []
        for msg in history[-10:]:
             role = "user" if msg["role"] == "user" else "model"
             contents.append(types.Content(
                 role=role,
                 parts=[types.Part.from_text(text=msg["content"])]
             ))
             
        contents.append(types.Content(
            role="user",
            parts=[types.Part.from_text(text=user_input)]
        ))
        
        # Generate response
        try:
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=contents,
                config=types.GenerateContentConfig(
                    system_instruction=system_prompt,
                    temperature=0.7,
                )
            )
            return response.text
        except Exception as e:
            return f"❌ Error generating response: {str(e)}"
    
    def export_conversation_to_docx(self, history: List[Dict[str, str]], title: str = "Conversación") -> bytes:
        """
        Export conversation to Word document
        
        Args:
            history: Conversation history
            title: Document title
            
        Returns:
            Bytes of the .docx file
        """
        if not HAS_DOCX:
            return b"Error: python-docx library not installed on server."

        doc = Document()
        
        # Title Style
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_CENTER
        
        doc.add_paragraph(f"Exportado desde Iglesia Irresistible OS")
        doc.add_paragraph("")
        
        for msg in history:
            role = "Usuario" if msg["role"] == "user" else "Asistente"
            
            # Role Header
            p = doc.add_paragraph()
            run = p.add_run(f"{role}:")
            run.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102) if msg["role"] == "user" else RGBColor(0, 102, 51)
            
            # Parse Content Line by Line
            content = msg["content"]
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # HEADERS
                if line.startswith('#'):
                    level = min(line.count('#'), 3)
                    clean_text = line.lstrip('#').strip()
                    doc.add_heading(clean_text, level=level)
                
                # BULLET POINTS
                elif line.startswith(('* ', '- ', '• ')):
                    clean_text = re.sub(r'^[\*\-•]\s+', '', line)
                    p = doc.add_paragraph(style='List Bullet')
                    self._add_formatted_run(p, clean_text)
                
                # NUMBERED LISTS (Basic detection)
                elif re.match(r'^\d+\.', line):
                    clean_text = re.sub(r'^\d+\.\s+', '', line)
                    p = doc.add_paragraph(style='List Number')
                    self._add_formatted_run(p, clean_text)
                    
                # STANDARD PARAGRAPH
                else:
                    p = doc.add_paragraph()
                    self._add_formatted_run(p, line)
            
            doc.add_paragraph("") # Spacer between messages
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_formatted_run(self, paragraph, text):
        """Helper to parse bold/italic within a paragraph run"""
        # Split by bold markers (**text**)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # Handle italics inside non-bold parts if needed, or just plain text
                # Simple update: just handle bold for now as it's the most common annoyance
                paragraph.add_run(part)

    def optimize_query(self, query: str) -> str:
        """
        Optimizes a search query by translating to English and expanding synonyms.
        """
        if not self.client:
            return query
            
        prompt = f"""
        You are a search expert. The user is searching for assets in a DAM (Digital Asset Management) system.
        The assets are mostly in English.
        Convert the following search term (which may be in Spanish) into an optimized search query string.
        
        RULES:
        1. Translate to English if in Spanish.
        2. Keep the original term.
        3. Add 1-2 key synonyms.
        4. Join with 'OR'.
        5. Return ONLY the query string, nothing else.
        
        Input: "{query}"
        """
        
        try:
            response = self.client.models.generate_content(
                model=self.model_name,
                contents=prompt
            )
            cleaned = response.text.strip().replace('"', '')
            print(f"🔍 Optimized Query: '{query}' -> '{cleaned}'")
            return cleaned
        except Exception as e:
            print(f"❌ Error optimizing query: {e}")
            return query
